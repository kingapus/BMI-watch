'''Determine Body Mass Index (BMI)'''

import sys
import speaking
import docx


#-----------------------------------------------------------------------
print(f"\n----------Program Start----------\n")
#-----------------------------------------------------------------------


print(f"Hello! Welcome...\n")
print(f"I can tell you your BMI if you are "
      "ready to give me some basic information about yourself. "
      "Additionally, If you give me your blood Pressure measurement,"
      "I can provide you with some essential advise on keeping healthy.\n")
print(f"\npress q to quit")

print(f"\nFirst, let me get to know you...\n")

f_name = input("What is your name: ")

#-----------------------------------------------------------------------
# Capture name and age
#-----------------------------------------------------------------------

if f_name.lower() == 'q':
    sys.exit()
while True: 
    try:
        age = int(input("\nHow old are you? "))
        break
    except ValueError:
        print(f"\nI do not recognize that as a valid age. Please enter"
              "your age.")
        continue
    
#-----------------------------------------------------------------------
# Set the loop for collecting weight and height measurements
#-----------------------------------------------------------------------

status0 = True
while status0:
    try:
        print(f"\nHey {f_name.title()}! would you like to see how useful"
              " I can be? Enter (1) or (2)\n")
        proceed = int(input(f"YES(1) or NO(2): "))
        options1, options2 = [1, 2,], ['q']
    ##try:
        if proceed in options1:
            if proceed == 1:
                print(f"\nLet's begin with your weight...")
                print(f"\nWhat unit would you like to use? ")
                print(f"Please press q to quit\n")
                print(f"\t--kg/m-- or --lbs/feet--")
                print(f"\t  (1)    or     (2) \n")
##            while True:
                try:
                    prompt = "Please enter (1) for kg or (2) for lbs: "
                    unit = int(input(prompt))
                    if unit in options1:
                        if unit == 1:
                            weight_unit = 'kg'
                            height_unit = 'm'
                            weight = float(input(f"\nPlease enter your weight in kilograms: "))
                            height = float(input(f"\nPlease enter your height in meters: "))
                            bmi = (weight)/(height**2)
                            print(f"\nHere are your details: ")
                            print(f"\n\tName: {f_name.title()}\n\tAge: {age}\n\tWeight: {weight}{weight_unit}\n\tHeight: {height}{height_unit}\n")
                            print(f"\t---------\n\tBMI: {round(bmi, 2)}kg/m2\n\t----------")
                            #break
                        if unit == 2:
                            weight_unit = 'lbs'
                            height_unit = 'feet'
                            weight = float(input(f"\nPlease enter your weight in pounds: "))
                            height = float(input(f"\nPlease enter your height in feet: "))
                            weight_in_kilos = weight/2.205
                            height_in_meters = height/3.281
                            bmi = (weight_in_kilos)/(height_in_meters**2)
                            print(f"Here are your details: ")
                            print(f"\n\tName: {f_name.title()}\n\tAge: {age}\n\tWeight: {weight}{weight_unit}\n\tHeight: {height}{height_unit}\n")
                            print(f"\t---------\n\tBMI: {round(bmi, 2)}kg/m2\n\t----------")
                            #break
                    else:
                        print(f"Incorrect input. Please try again.\n")
                        continue
                except:
                    print(f"Oops, that is not a valid input. Please re-try.")
            elif proceed == 2:
                print(f"Sad to see you leave :(  Goodbye and have a healthy"
                      " day.")
                status0 = False


        elif proceed not in options1:
            if proceed > 2:
                print(f"Invalid input.")

    except:
        print(f"\nInvalid input.\n")
    break

if proceed == 2:
    sys.exit()

#-----------------------------------------------------------------------
# Request of blood pressure data
#-----------------------------------------------------------------------


status1 = True
while status1:
    try:
        print(f"\nWould you like to add your blood pressure? "
              "It would help me give you a comprehensive advice")
        options = [1, 2]
        bp_option = int(input(f"\nYes(1) or No(2): "))

        if bp_option in options:
            if bp_option == 1:
                print("Ok\n")
                sbp = int(input("Please enter your systolic blood pressure: "))
                dbp = int(input("Please enter your diastolic blood pressure: "))
                break
            if bp_option == 2:
                print(f"\nOk")
                sbp = 0
                dbp = 0
                status1 = False             


    except:
        print(f"\nPlease select or enter a valid response")
        
#-----------------------------------------------------------------------
# BMI and Blood Pressure notes/advice categorised
#-----------------------------------------------------------------------

file = docx.Document("BMI and BP.docx")

#--------------------------------------------------
# Underweight
#--------------------------------------------------

underweight_no_bp = file.paragraphs[1].text
underweight_hypertensive = file.paragraphs[3].text
underweight_normotensive = file.paragraphs[5].text
what_to_do = file.paragraphs[25].text

#--------------------------------------------------
# Normal
#--------------------------------------------------

normal_no_bp = file.paragraphs[7].text
normal_hypertensive = file.paragraphs[9].text
normal_normotensive = file.paragraphs[11].text
what_to_do = file.paragraphs[25].text

#--------------------------------------------------
# Overweight
#--------------------------------------------------

overweight_no_bp = file.paragraphs[13].text
overweight_hypertensive = file.paragraphs[15].text
overweight_normotensive = file.paragraphs[17].text
what_to_do = file.paragraphs[25].text

#--------------------------------------------------
# Obese
#--------------------------------------------------

obese_no_bp = file.paragraphs[19].text
obese_hypertensive = file.paragraphs[21].text
obese_normotensive = file.paragraphs[23].text
what_to_do = file.paragraphs[25].text

#--------------------------------------------------
#--------------------------------------------------

#-----------------------------------------------------------------------
# Display and read health advice based on BMI and Blood Pressure
#-----------------------------------------------------------------------


if bmi < 18.5:
    print(f"\nHi {f_name.title()}...\n")
    print(f"\nYour BMI is {round(bmi, 2)} kg/m2; and your blood pressure is {sbp}/{dbp} mmHg\n")
    if (sbp + dbp) < 230:
        speaking.speak(underweight_normotensive, what_to_do)
    elif (sbp + dbp) == 0:
        speaking.speak(underweight_no_bp, what_to_do)
    else:
        speaking.speak(underweight_hypertensive, what_to_do)
elif bmi < 24.9:
    print(f"\nHi {f_name.title()}...\n")
    print(f"\nYour BMI is {round(bmi, 2)} kg/m2; and your blood pressure is {sbp}/{dbp} mmHg\n")
    if (sbp + dbp) < 230:
        speaking.speak(normal_normotensive, what_to_do)
    elif (sbp + dbp) == 0:
        speaking.speak(normal_no_bp, what_to_do)
    else:
        speaking.speak(normal_hypertensive, what_to_do)
elif bmi < 30:
    print(f"\nHi {f_name.title()}...\n")
    print(f"\nYour BMI is {round(bmi, 2)} kg/m2; and your blood pressure is {sbp}/{dbp} mmHg\n")    
    if (sbp + dbp) < 230:
        speaking.speak(overweight_normotensive, what_to_do)
    elif (sbp + dbp) == 0:
        speaking.speak(overweight_no_bp, what_to_do)
    else:
        speaking.speak(overweight_hypertensive, what_to_do)
elif bmi >= 30:
    print(f"\nHi {f_name.title()}...\n")
    print(f"\nYour BMI is {round(bmi, 2)} kg/m2; and your blood pressure is {sbp}/{dbp} mmHg\n")
    if (sbp + dbp) < 230:
        speaking.speak(obese_normotensive, what_to_do)
    elif (sbp + dbp) == 0:
        speaking.speak(obese_no_bp, what_to_do)
    else:
        speaking.speak(obese_hypertensive, what_to_do)


#-----------------------------------------------------------------------
print(f"\n----------Program End----------\n")
#-----------------------------------------------------------------------

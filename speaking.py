import pyttsx3
import docx

file = docx.Document("BMI and BP.docx")

#--------------------------------------------------
# Underweight
#--------------------------------------------------

underweight_no_bp = file.paragraphs[1].text
underweight_hypertensive = file.paragraphs[3].text
underweight_normotensive = file.paragraphs[5].text
what_to_do = file.paragraphs[25].text

#--------------------------------------------------
# Normal
#--------------------------------------------------

normal_no_bp = file.paragraphs[7].text
normal_hypertensive = file.paragraphs[9].text
normal_normotensive = file.paragraphs[11].text
what_to_do = file.paragraphs[25].text

#--------------------------------------------------
# Overweight
#--------------------------------------------------

overweight_no_bp = file.paragraphs[13].text
overweight_hypertensive = file.paragraphs[15].text
overweight_normotensive = file.paragraphs[17].text
what_to_do = file.paragraphs[25].text

#--------------------------------------------------
# Obese
#--------------------------------------------------

obese_no_bp = file.paragraphs[19].text
obese_hypertensive = file.paragraphs[21].text
obese_normotensive = file.paragraphs[23].text
what_to_do = file.paragraphs[25].text

#--------------------------------------------------
# speaking function
#--------------------------------------------------

def speak (category, action):
    print(category)
    print("\n")
    print(action)
    engine = pyttsx3.init()
    rate = engine.getProperty("rate")
    engine.setProperty("rate", 150)
    engine.say("Here are your results and some health tips specific to your results")
    engine.say(category)
    engine.say("Here is what you should keep doing")
    engine.say(action)
    engine.say("thank you!")
    engine.runAndWait()
    

def welcome():
    engine = pyttsx3.init()
    rate = engine.getProperty("rate")
    engine.setProperty("rate", 150)
    engine.say("Program Start")

